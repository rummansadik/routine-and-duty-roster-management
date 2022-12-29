import datetime
import io
import os
from time import strftime

from account.models import *
from django.conf import settings
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Mm

from dashboard.models import *


def get_data(exams: Exam):
    exams_data = []
    for exam in exams:
        data = []
        t = exam.exam_time.strftime("%I:%M %P")
        data.append(str(exam.exam_date) + "\n" + t)
        data.append(exam.course.code)
        data.append(exam.supervisor.get_name)
        names = list(exam.examiners.all())
        names = list(map(lambda x: x.get_name, names))
        data.append("\n".join(names))
        exams_data.append(data)
    return exams_data


def add_text(row, col, msg, flag=False):
    row[col].text = msg
    p = row[col].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Inches(.2)
    if flag:
        p.paragraph_format.space_after = Inches(.2)


def add_header_section(doc, routine: Routine, level, semester):
    data = [
        'Duty Roaster List',
        f"B'Sc in {routine.department}",
        "Semester Final Examination",
        f"Level: {level}, Semester: {semester}",
    ]

    title = doc.add_paragraph("\n".join(data))
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER


def get_level(exam: Exam):
    level = exam.course.level
    semester = exam.course.semester
    return level, semester


def roaster(exams: Exam, routine: Routine):

    doc = Document()

    level, semester = get_level(exams[0])
    add_header_section(doc, routine, level, semester)

    table = doc.add_table(rows=1, cols=4, style='Table Grid')

    table.columns[0].width = Inches(1.2)
    table.columns[1].width = Inches(1.2)
    table.columns[2].width = Inches(2.0)
    table.columns[3].width = Inches(2.0)

    row = table.rows[0].cells
    add_text(row, 0, "Date")
    add_text(row, 1, "Course Code")
    add_text(row, 2, "Supervisor")
    add_text(row, 3, "Examiners")
    # add_text(row, 4, "Assistance")

    data = get_data(exams)
    for date, code, supervisor, examiners in data:
        row = table.add_row().cells
        add_text(row, 0, date)
        add_text(row, 1, code)
        add_text(row, 2, supervisor)
        add_text(row, 3, examiners, True)
        # add_text(row, 4, assistance)

    base_dir = settings.BASE_DIR
    file = os.path.join(base_dir, f"Sample/Roaster/{routine.name}.docx")
    doc.save(file)
